import os
import base64
import json
import pandas as pd
from openai import OpenAI
from pdf2image import convert_from_path
from PIL import Image
import io
from docx import Document
from docx.shared import Inches

# 配置参数
MODEL_ID = "Qwen/Qwen2.5-VL-72B-Instruct"

def pdf_to_images(pdf_path, output_dir=None):
    """将PDF转换为图片列表"""
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 转换PDF为图片
    images = convert_from_path(pdf_path)
    
    # 如果指定了输出目录，则保存图片
    if output_dir:
        image_paths = []
        pdf_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        for i, image in enumerate(images):
            image_name = f"{pdf_filename}_page_{i+1}.jpg"
            image_path = os.path.join(output_dir, image_name)
            image.save(image_path, "JPEG")
            image_paths.append(image_path)
        return image_paths
    
    return images

def image_to_base64(image):
    """将PIL Image对象或图片路径转换为Base64编码"""
    if isinstance(image, str):
        if not os.path.exists(image):
            raise FileNotFoundError(f"图片文件不存在: {image}")
        with open(image, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode('utf-8')
    
    # 处理PIL Image对象
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def analyze_image_with_model(client, image, model_id=MODEL_ID):
    """使用多模态模型分析图片内容"""
    image_base64 = image_to_base64(image)
    
    response = client.chat.completions.create(
        model=model_id,
        messages=[{
            'role': 'user',
            'content': [{
                'type': 'text',
                'text': '详细描述这幅图中的所有内容，包括文本、图表、图像元素等，直接介绍具体内容。如果存在表格，额外将表格输出为markdown的格式。要求表头准确清晰，注意表头的层级结构内容完整且对应无误，请确保表格内容准确对应到相应列，表格内容不要翻译。',
            }, {
                'type': 'image_url',
                'image_url': {
                    'url': f'data:image/jpeg;base64,{image_base64}'
                },
            }],
        }]
    )
    
    return response.choices[0].message.content

def json_to_docx(results, output_dir, pdf_name):
    """将JSON格式的结果转换为Word文档"""
    doc = Document()
    doc.add_heading(f'{pdf_name} 内容提取结果', 0)
    
    for result in results:
        page_num = result['page_number']
        doc.add_heading(f'{pdf_name} 第 {page_num} 页', level=1)
        
        # 添加图片（如果有）
        if 'image_name' in result and result['image_name']:
            image_path = os.path.join(output_dir, "images", result['image_name'])
            if os.path.exists(image_path):
                try:
                    doc.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[无法添加图片: {str(e)}]")
            else:
                doc.add_paragraph(f"[图片文件不存在: {image_path}]")
        
        # 添加内容
        if 'error' in result:
            doc.add_paragraph(f"处理错误: {result['error']}")
        else:
            doc.add_paragraph(result['content'])
        
        # 添加分页符（除了最后一页）
        if page_num < len(results):
            doc.add_page_break()
    
    # 保存文档
    output_path = os.path.join(output_dir, f"{os.path.splitext(pdf_name)[0]}_extracted.docx")
    doc.save(output_path)
    print(f"Word文档已保存至: {output_path}")

def process_pdf(pdf_path, output_dir=None, progress=None):
    """处理单个PDF文件，将其转换为图片并提取文本内容"""
    if progress:
        progress['currentFile'] = os.path.basename(pdf_path)
        progress['status'] = '正在处理'
    
    # 初始化OpenAI客户端
    client = OpenAI(
        base_url='https://api-inference.modelscope.cn/v1/',
        api_key='your_api_key',
    )
    
    # 创建图片输出目录
    images_output_dir = None
    if output_dir:
        images_output_dir = os.path.join(output_dir, "images")
        os.makedirs(images_output_dir, exist_ok=True)
    
    pdf_filename = os.path.basename(pdf_path)
    
    # 将PDF转换为图片
    print(f"正在将PDF转换为图片: {pdf_path}")
    images = pdf_to_images(pdf_path, images_output_dir)
    
    # 处理每一页图片
    results = []
    total_pages = len(images)
    if progress:
        progress['total'] += total_pages
    
    for page_num, image in enumerate(images, 1):
        print(f"正在处理第 {page_num} 页...")
        
        try:
            image_name = f"{os.path.splitext(pdf_filename)[0]}_page_{page_num}.jpg" if images_output_dir else None
            page_content = analyze_image_with_model(client, image)
            
            result_entry = {
                "page_number": page_num,
                "image_name": image_name,
                "content": page_content,
                "source_pdf": pdf_filename,
                "timestamp": str(pd.Timestamp.now())
            }
            
            results.append(result_entry)
        except Exception as e:
            print(f"处理第 {page_num} 页时出错: {str(e)}")
            results.append({
                "page_number": page_num,
                "image_name": f"{os.path.splitext(pdf_filename)[0]}_page_{page_num}.jpg" if images_output_dir else None,
                "error": str(e),
                "source_pdf": pdf_filename,
                "timestamp": str(pd.Timestamp.now())
            })

        if progress:
            progress['current'] += 1
    
    # 保存结果
    if output_dir:
        # 保存JSON文件
        json_output_path = os.path.join(output_dir, f"{os.path.splitext(pdf_filename)[0]}_extracted.json")
        with open(json_output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        
        print(f"JSON结果已保存至: {json_output_path}")
        
        # 转换为Word文档
        json_to_docx(results, output_dir, pdf_filename)
    
    return results 